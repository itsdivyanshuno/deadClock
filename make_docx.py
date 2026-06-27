#!/usr/bin/env python3
"""
Generate the hackathon submission document (HTML, Markdown, DOCX).

Outputs:
  notes/submission-doc.html
  notes/SUBMISSION_DOC.md
  notes/submission-doc.docx   (if python-docx is installed)
"""

from __future__ import annotations

import pathlib
from textwrap import dedent, indent


REPO_ROOT = pathlib.Path(__file__).parent
NOTES = REPO_ROOT / "notes"
NOTES.mkdir(exist_ok=True)


# ── Content ────────────────────────────────────────────────────────────────────
TITLE = "deadClock — The Last-Minute Life Saver"
TAGLINE = "An AI productivity companion that plans, prioritises, and helps you beat every deadline."
TECH_STACK = [
    ("Framework", "Next.js 14 (App Router)"),
    ("Language", "TypeScript + React"),
    ("Styling", "Tailwind CSS v3"),
    ("AI Engine", "OpenAI API (GPT-4o / GPT-4o Mini)"),
    ("State Management", "Zustand"),
    ("Charts", "Recharts"),
    ("Rich Text", "Tiptap"),
    ("Persistence", "LocalStorage / sessionStorage"),
    ("Routing", "Next.js App Router (file-based)"),
    ("Animations", "Framer Motion"),
    ("Icons", "Lucide React"),
]
FEATURES = [
    "AI-native task management – chat naturally with an assistant that creates, prioritises, and completes tasks (OpenAI function calling).",
    "Goal tracking with milestones – set long-term goals and break them into proportional weekly tasks automatically.",
    "At-risk deadline detection – the AI identifies tasks within a 24-hour risk window and suggests realistic reschedules.",
    "Proactive workload assistant – based on your current load, overdue items, and goal progress, the assistant surfaces what you should focus on right now.",
    "Clean three-pane interface – sidebar overview, main chat surface, and task/goal views that stay in sync automatically.",
    "Local-first persistence – all data stays in your browser (Zustand + LocalStorage). No accounts, no cloud sync.",
]
GEMINI_TOOLS = [
    ("add_task", "Create a task with title, deadline, priority, category, and subtasks."),
    ("prioritize_tasks", "Re-sort pending tasks by urgency and impact."),
    ("complete_task", "Mark a task done by ID."),
    ("add_goal", "Create a long-term goal with milestones."),
    ("suggest_schedule", "Generate a time-blocked daily plan."),
    ("get_reminders", "Surface urgent / overdue tasks."),
    ("suggest_proactive_actions", "Analyse workload and surface proactive suggestions."),
    ("break_down_goal", "Split a goal into proportional weekly tasks."),
    ("reschedule_at_risk_tasks", "Extend deadlines with a safety buffer."),
]


def build_md() -> str:
    """Return the submission document as Markdown."""
    md_parts = [
        f"# {TITLE}\n",
        f"> {TAGLINE}\n",
        "---\n",
        "## Tech Stack\n",
        "| Layer | Choice | Role |",
        "|-------|--------|------|",
    ]
    for layer, detail in TECH_STACK:
        role = detail  # detail already includes role in the tuple
        md_parts.append(f"| {layer} | {detail} | — |")
    md_parts += [
        "",
        "---",
        "## Core Features",
        "",
    ]
    for feat in FEATURES:
        md_parts.append(f"- {feat}")
    md_parts += [
        "",
        "---",
        "## How It Works",
        "",
        "```",
        "User types message → POST /api/chat → OpenAI agent (lib/agent.ts)",
        "  ├─ Loads state from LocalStorage",
        "  ├─ Sends system prompt + chat history → GPT-4o",
        "  ├─ GPT-4o may call tools: add_task, prioritise_tasks, …",
        "  ├─ Tool results mutate in-memory state",
        "  ├─ GPT-4o produces a follow-up text response",
        "  ├─ State is persisted to LocalStorage (Zustand)",
        "  └─ Returns { response, tasks[], goals[] } → Client updates state",
        "```",
        "",
        "The client never mutates the server-side store directly. Every change flows through the OpenAI agent via defined tool contracts.",
        "",
        "---",
        "## OpenAI Tools (Function Calling)",
        "",
        "| Tool | Purpose |",
        "|------|---------|",
    ]
    for name, desc in GEMINI_TOOLS:
        md_parts.append(f"| `{name}` | {desc} |")
    md_parts += [
        "",
        "---",
        "## Agentic Depth",
        "",
        "- **9 OpenAI function-calling tools** — the AI has full CRUD + planning authority.",
        "- **Proactive workload analysis** — surfaces overdue items and load imbalances unprompted.",
        "- **Autonomous goal breakdown** — splits long-term goals into proportional weekly tasks.",
        "- **Risk-aware rescheduling** — tasks due within 24 h get a safety-buffer extension.",
        "",
        "---",
        "## UI / UX",
        "",
        "- **Three-pane desktop layout** — sidebar (overview + quick-nav), main canvas (chat / tasks / goals).",
        "- **Smooth animations** — Framer Motion transitions on view switches, message bubbles, and card hover states.",
        "- **Custom minimap** — visual task-progress renderer in the sidebar.",
        "- **Slash commands** — type `/help`, `/schedule`, `/goals`, `/clear` in chat for fast actions.",
        "",
        "---",
        "## Project Structure",
        "",
        "```",
        "app/",
        "├── layout.tsx       ← Root HTML wrapper + global providers",
        "├── page.tsx         ← Main app shell (sidebar + canvas)",
        "├── globals.css      ← Tailwind + design tokens",
        "├── api/",
        "│   └── openai/",
        "│       └── route.ts ← Chat endpoint (Node.js runtime)",
        "├── chat/",
        "│   └── page.tsx     ← Dedicated full-screen chat view",
        "└── components/",
        "    ├── Chat/…       ← Message bubbles, input bar, slash commands",
        "    ├── TaskBoard/…  ← Kanban-style task management",
        "    ├── GoalTracker/ ← Timeline + milestone progress",
        "    ├── Analytics/   ← Charts (Recharts)",
        "    └── Settings/    ← Preferences pane",
        "",
        "lib/",
        "├── agent.ts      ← OpenAI agent: tools, orchestrator, helpers",
        "├── store.ts      ← Zustand store + LocalStorage sync",
        "└── utils.ts      ← Date helpers, ID generation",
        "",
        "notes/",
        "└── SUBMISSION_DOC.md ← This document",
        "```",
        "",
        "---",
        "",
        f"**Built with ❤️ for Vibe2Ship Hackathon 2026**  ",
        f"Repository: https://github.com/itsdivyanshuno/vibe2ship-lastminute-lifesaver",
    ]
    return "\n".join(md_parts)


def build_html(md_text: str) -> str:
    """Return a self-contained HTML preview of the Markdown document."""
    return dedent(f"""\
    <!DOCTYPE html>
    <html lang="en">
    <head>
      <meta charset="utf-8" />
      <meta name="viewport" content="width=device-width, initial-scale=1" />
      <title>{TITLE} — Submission Preview</title>
      <style>
        /* ── Reset ────────────────────────────────────────────── */
        *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
        /* ── Base ─────────────────────────────────────────────── */
        body {{
          font-family: "Inter", system-ui, -apple-system, "Segoe UI", Roboto,
                       "Helvetica Neue", Arial, sans-serif;
          font-size: 15px;
          line-height: 1.7;
          color: #1c1917;
          background: #fafaf9;
          padding: 2rem 1rem;
        }}
        .container {{
          max-width: 820px;
          margin: 0 auto;
        }}
        /* ── Typography ───────────────────────────────────────── */
        h1 {{ font-size: 2rem; font-weight: 800; letter-spacing: -0.03em; margin-bottom: .5rem; }}
        h2 {{ font-size: 1.35rem; font-weight: 700; letter-spacing: -0.02em;
              margin: 2.5rem 0 1rem; padding-bottom: .5rem;
              border-bottom: 2px solid #e7e5e4; }}
        p {{ margin-bottom: 1rem; }}
        blockquote {{ border-left: 4px solid #ea580c; padding: .5rem 1rem;
                      background: #fff7ed; border-radius: 0 8px 8px 0;
                      font-style: italic; margin-bottom: 1.5rem; }}
        code {{ font-family: "JetBrains Mono", "Fira Code", monospace;
                font-size: .85em; background: #f1f5f9; padding: .15em .4em;
                border-radius: 4px; }}
        pre {{ background: #1e293b; color: #e2e8f0; padding: 1.25rem;
               border-radius: 12px; overflow-x: auto; margin-bottom: 1.5rem;
               font-size: .82rem; line-height: 1.6; }}
        pre code {{ background: none; padding: 0; color: inherit; }}
        table {{ width: 100%; border-collapse: collapse; margin-bottom: 1.5rem; }}
        th {{ text-align: left; font-size: .75rem; text-transform: uppercase;
              letter-spacing: .06em; color: #57534e; padding: .5rem .75rem;
              border-bottom: 2px solid #e7e5e4; }}
        td {{ padding: .55rem .75rem; border-bottom: 1px solid #f5f5f4;
              font-size: .9rem; vertical-align: top; }}
        tr:last-child td {{ border-bottom: none; }}
        ul {{ padding-left: 1.5rem; margin-bottom: 1rem; }}
        li {{ margin-bottom: .4rem; }}
        hr {{ border: none; border-top: 1px solid #e7e5e4; margin: 2rem 0; }}
        .meta {{ color: #57534e; font-size: .9rem; }}
        a {{ color: #1e293b; text-decoration: underline; text-underline-offset: 2px; }}
        a:hover {{ color: #ea580c; }}
        /* ── Header ────────────────────────────────────────────── */
        header {{ margin-bottom: 3rem; }}
        .slug {{ display: inline-block; font-size: .7rem; font-weight: 600;
                text-transform: uppercase; letter-spacing: .08em;
                background: #1e293b; color: white; padding: .3rem .7rem;
                border-radius: 999px; margin-bottom: 1rem; }}
        .tech-pills {{ display: flex; flex-wrap: wrap; gap: .5rem; margin-top: 1rem; }}
        .pill {{ font-size: .78rem; font-weight: 500; padding: .3rem .75rem;
                border-radius: 999px; background: #f1f5f9; color: #1e293b;
                border: 1px solid #e7e5e4; }}
        .footer {{ margin-top: 4rem; padding-top: 1.5rem; border-top: 1px solid #e7e5e4;
                   font-size: .85rem; color: #57534e; }}
        @media (max-width: 600px) {{
          body {{ padding: 1rem .75rem; }}
          h1 {{ font-size: 1.5rem; }}
        }}
      </style>
    </head>
    <body>
      <div class="container">
        <header>
          <span class="slug">Hackathon Submission</span>
          <h1>{TITLE}</h1>
          <p class="meta">{TAGLINE}</p>
          <div class="tech-pills">
            {"".join(f'<span class="pill">{layer}: {detail}</span>' for layer, detail in TECH_STACK)}
          </div>
        </header>

        <hr />

        <!-- Features -->
        <h2>🚀 Core Features</h2>
        <ul>
          {"".join(f"<li>{f}</li>" for f in FEATURES)}
        </ul>

        <hr />

        <!-- How it works -->
        <h2>⚙️ How It Works</h2>
        <p>deadClock isn't a todo list you have to maintain. It's an AI agent that has a <em>conversation</em> with you:</p>
        <pre><code>{dedent("""\
        User types message
          → POST /api/chat
          → OpenAI agent orchestrator (lib/agent.ts)
          → Load state from LocalStorage / Zustand
          → Send system prompt + chat history → GPT-4o
          → GPT-4o may call tools (add_task, prioritise_tasks, …)
          → Tool results mutate in-memory state
          → GPT-4o produces natural-language follow-up
          → State persisted to LocalStorage
          → Returns { response, tasks[], goals[] }
          → Client updates UI
        """).strip()}</code></pre>

        <hr />

        <!-- Tools -->
        <h2>🤖 OpenAI Tools (Function Calling)</h2>
        <table>
          <tr><th>Tool</th><th>What it does</th></tr>
          {"".join(f"<tr><td><code>{name}</code></td><td>{desc}</td></tr>" for name, desc in GEMINI_TOOLS)}
        </table>

        <hr />

        <!-- Agentic features -->
        <h2>🧠 Agentic Depth</h2>
        <ul>
          <li><strong>9 OpenAI function-calling tools</strong> — the AI has full CRUD + planning authority over your tasks and goals.</li>
          <li><strong>Proactive workload analysis</strong> — surfaces overdue items, load imbalances, and lagging goals unprompted.</li>
          <li><strong>Autonomous goal breakdown</strong> — splits long-term goals into proportional weekly tasks based on your available hours.</li>
          <li><strong>Risk-aware rescheduling</strong> — tasks due within 24 hours automatically get a safety-buffer deadline extension.</li>
        </ul>

        <hr />

        <!-- UI -->
        <h2>🎨 UI / UX</h2>
        <ul>
          <li><strong>Three-pane desktop layout</strong> — sidebar (overview + nav), main canvas (chat / tasks / goals panels).</li>
          <li><strong>Smooth animations</strong> — Framer Motion transitions on view switches, message bubbles, and card hover states.</li>
          <li><strong>Custom minimap</strong> — minimap-style SVG based task view on the canvas.</li>
          <li><strong>Slash commands</strong> — type <code>/help</code>, <code>/schedule</code>, <code>/goals</code>, <code>/clear</code> in chat for fast actions.</li>
        </ul>

        <hr />

        <!-- Project structure -->
        <h2>📁 Project Structure</h2>
        <pre><code>{dedent("""\
        app/
        ├── layout.tsx       ← Root HTML wrapper + global providers
        ├── page.tsx         ← Main app shell (sidebar + canvas)
        ├── globals.css      ← Tailwind + design tokens
        ├── api/
        │   └── chat/
        │       └── route.ts ← Chat endpoint (Node.js runtime)
        ├── chat/
        │   └── page.tsx     ← Dedicated full-screen chat view
        └── components/
             ├── Sidebar/…   ← Nav, goal cards, task overview
             ├── Chat/…      ← Message bubbles, input bar, slash commands
             ├── TaskBoard/  ← Kanban-style task management
             ├── GoalTracker/← Timeline + milestone progress
             └── Settings/   ← Preferences pane

        lib/
        ├── agent.ts      ← OpenAI agent: tools, orchestrator, helpers
        ├── store.ts      ← Zustand store + sessionStorage sync
        └── utils.ts      ← Date helpers, ID generation

        notes/
        └── SUBMISSION_DOC.md ← This document
        """).strip()}</code></pre>

        <hr />

        <!-- Architecture notes -->
        <h2>📐 Architecture Notes</h2>
        <ul>
          <li><strong>OpenAI-first AI layer.</strong> `chat()` in `lib/agent.ts` is the single entry point — it loads state, calls GPT-4o, dispatches function calls, and returns structured tool results.</li>
          <li><strong>Client is a thin orchestrator.</strong> `app/page.tsx` holds UI state and delegates to a single fetch against <code>/api/chat</code>.</li>
          <li><strong>Local-first.</strong> No accounts, no cloud sync, no cookies. State is in Zustand + sessionStorage.</li>
          <li><strong>Runtime-aware routing.</strong> The API route declares <code>export const runtime = "nodejs"</code> so the OpenAI SDK runs on the server, not the Edge.</li>
        </ul>

        <hr />

        <!-- Getting started -->
        <h2>🔧 Getting Started</h2>
        <pre><code>{dedent("""\
        git clone https://github.com/itsdivyanshuno/vibe2ship-lastminute-lifesaver.git
        cd vibe2ship-lastminute-lifesaver
        npm install
        cp .env.local.example .env.local   # paste your OpenAI key
        npm run dev
        """).strip()}</code></pre>
        <p>Open <a href="http://localhost:3000">http://localhost:3000</a> and tell the AI about your day.</p>

        <hr />

        <h2>🏆 What's Inside</h2>
        <table>
          <tr><th>Area</th><th>Highlights</th></tr>
          <tr><td>Agentic depth</td><td>9 OpenAI tools, proactive suggestions, autonomous goal breakdown, risk-aware rescheduling</td></tr>
          <tr><td>Local-first</td><td>Zustand + sessionStorage, no accounts, no cloud, no cookies</td></tr>
          <tr><td>Performance</td><td>Turbopack (Next 14), client-side routing, Framer Motion GPU-accelerated</td></tr>
          <tr><td>Code quality</td><td>JSDoc on every public function, typed interfaces, explicit error contracts</td></tr>
        </table>

        <div class="footer">
          <p><strong>Built for the Vibe2Ship Hackathon</strong> — June 22–29, 2026</p>
          <p>
            <a href="https://github.com/itsdivyanshuno/vibe2ship-lastminute-lifesaver">
              github.com/itsdivyanshuno/vibe2ship-lastminute-lifesaver
            </a>
          </p>
        </div>
      </div>
    </body>
    </html>
    """)


def build_docx(md_text: str) -> None:
    """Create a .docx from the Markdown using python-docx (optional dependency)."""
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt, RGBColor  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    except ImportError:
        print("[make_docx] python-docx is not installed — skipping .docx output.")
        print("  Install with: pip install python-docx")
        return

    doc = Document()

    # Set default font for the document
    doc.styles["Normal"].font.name = "Inter"
    doc.styles["Normal"].font.size = Pt(11)
    doc.styles["Normal"].font.color.rgb = RGBColor(0x1C, 0x19, 0x17)

    # Title
    title = doc.add_heading(TITLE, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    sub = doc.add_paragraph(TAGLINE)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True

    doc.add_paragraph("")  # spacer

    # Tech stack
    doc.add_heading("Tech Stack", level=1)
    for layer, detail in TECH_STACK:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{layer}: ").bold = True
        p.add_run(detail)

    # Features
    doc.add_heading("Core Features", level=1)
    for feat in FEATURES:
        doc.add_paragraph(feat, style="List Bullet")

    # How it works
    doc.add_heading("How It Works", level=1)
    for line in [
        "1. User types a message in the chat UI.",
        "2. The message is POSTed to /api/chat.",
        "3. The OpenAI agent loads state from LocalStorage.",
        "4. GPT-4o processes the message with tool definitions.",
        "5. Tool calls mutate in-memory state (tasks, goals).",
        "6. GPT-4o produces a natural-language follow-up.",
        "7. State is persisted back to LocalStorage.",
        "8. Client receives { response, tasks[], goals[] } and updates.",
    ]:
        doc.add_paragraph(line, style="List Number")

    # Tools table
    doc.add_heading("OpenAI Tools", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    hdr = table.rows[0].cells
    hdr[0].text = "Tool"
    hdr[1].text = "What it does"
    for name, desc in GEMINI_TOOLS:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = desc

    # Agentic depth
    doc.add_heading("Agentic Depth", level=1)
    for item in [
        "9 OpenAI function-calling tools — the AI has full CRUD + planning authority.",
        "Proactive workload analysis — surfaces overdue and at-risk items.",
        "Autonomous goal breakdown — splits goals into proportional weekly tasks.",
        "Risk-aware rescheduling — extends deadlines within a configurable buffer.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # UI / UX
    doc.add_heading("UI / UX", level=1)
    for item in [
        "Three-pane desktop layout — sidebar (tasks/goals) + main canvas (minimap + chat).",
        "Smooth Framer Motion transitions and micro-interactions.",
        "Custom SVG minimap for task-progress visualization.",
        "Slash commands for fast actions.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # Project structure
    doc.add_heading("Project Structure", level=1)
    doc.add_paragraph(
        "app/\n"
        "├── layout.tsx       ← Root HTML wrapper + global providers\n"
        "├── page.tsx         ← Main app shell (sidebar + canvas)\n"
        "├── globals.css      ← Tailwind + design tokens\n"
        "├── api/\n"
        "│   └── chat/\n"
        "│       └── route.ts ← Chat endpoint (Node.js runtime)\n"
        "├── chat/\n"
        "│   └── page.tsx     ← Full-screen chat view\n"
        "└── components/       ← Sidebar, Chat, TaskBoard, GoalTracker, Settings\n\n"
        "lib/\n"
        "├── agent.ts      ← OpenAI agent\n"
        "├── store.ts      ← Zustand store + sessionStorage\n"
        "└── utils.ts      ← Date helpers, ID generation\n\n"
        "notes/\n"
        "└── SUBMISSION_DOC.md ← This document"
    )

    # Architecture notes
    doc.add_heading("Architecture Notes", level=1)
    for item in [
        "OpenAI-first AI layer — chat() in lib/agent.ts is the single entry point.",
        "Client is a thin orchestrator — app/page.tsx delegates to a single /api/chat fetch.",
        "Local-first — no accounts, no cloud sync, no cookies.",
        "Runtime-aware — API route declares Node.js runtime for the OpenAI SDK.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph(
        "Built for the Vibe2Ship Hackathon — June 22–29, 2026\n"
        "github.com/itsdivyanshuno/vibe2ship-lastminute-lifesaver"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    out = NOTES / "submission-doc.docx"
    doc.save(str(out))
    print(f"[make_docx] Written: {out}")


def main() -> None:
    md = build_md()
    md_path = NOTES / "SUBMISSION_DOC.md"
    md_path.write_text(md, encoding="utf-8")
    print(f"[make_docx] Written: {md_path}")

    html = build_html(md)
    html_path = NOTES / "submission-doc.html"
    html_path.write_text(html, encoding="utf-8")
    print(f"[make_docx] Written: {html_path}")

    build_docx(md)


if __name__ == "__main__":
    main()
